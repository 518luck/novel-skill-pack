#!/usr/bin/env python3
"""整合短剧素材并导出 Markdown / DOCX。

默认从当前项目读取：.drama-state.json、characters.md、episodes/。
也可以通过 DRAMA_PROJECT_ROOT 指定项目目录。

示例（由 Codex 的内置 Python 环境执行）：
    python build_export_docx.py
"""

from datetime import date
import json
import os
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


def find_project_root():
    configured = os.environ.get('DRAMA_PROJECT_ROOT')
    candidates = []
    if configured:
        candidates.append(Path(configured).expanduser().resolve())
    candidates.append(Path.cwd().resolve())
    candidates.append(Path(__file__).resolve().parents[4])
    for candidate in candidates:
        if (candidate / '.drama-state.json').exists():
            return candidate
    raise FileNotFoundError('找不到项目根目录，请设置 DRAMA_PROJECT_ROOT。')


ROOT = find_project_root()
EXPORT = ROOT / 'export'
STATE_PATH = ROOT / '.drama-state.json'
CHAR_PATH = ROOT / 'characters.md'
EPISODES = ROOT / 'episodes'

FONT = '宋体'
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(122, 90, 0)
GRAY = RGBColor(85, 85, 85)


def set_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for key in ('ascii', 'hAnsi', 'eastAsia'):
        rfonts.set(qn(f'w:{key}'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name=FONT, size=11, color=RGBColor(0, 0, 0), bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for key in ('ascii', 'hAnsi', 'eastAsia'):
        rfonts.set(qn(f'w:{key}'), name)


def set_spacing(style, before=0, after=8, line=1.333, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.alignment = alignment


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, end])
    set_font(run, size=9, color=GRAY)


def add_para(doc, text='', style='Normal', align=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def parse_roles():
    text = CHAR_PATH.read_text(encoding='utf-8')
    matches = list(re.finditer(r'^### (.+)$', text, re.M))
    roles = []
    seen = set()

    def find_field(section, kind):
        if kind == 'outer':
            pattern = r'^\s*(?:-\s*)?\*\*(外貌特征|外貌与记忆点|外貌|外形|形象)：\*\*\s*(.+)$'
        else:
            pattern = r'^\s*(?:-\s*)?\*\*(性格描述|性格关键词|性格|人物性格)：\*\*\s*(.+)$'
        match = re.search(pattern, section, re.M)
        return match.group(2).strip() if match else ''

    for index, match in enumerate(matches):
        heading = match.group(1).strip()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section = text[match.end():end]
        name = identity = ''
        if '｜' in heading:
            parts = [part.strip() for part in heading.split('｜')]
            name = re.sub(r'^[1-9]\d*\.\s*', '', parts[0])
            identity = parts[1] if len(parts) > 1 else '重要角色'
        elif re.match(r'^.+?[（(].+?[）)]$', heading):
            parsed = re.match(r'^(.+?)[（(](.+?)[）)]$', heading)
            name = parsed.group(1).strip()
            identity = parsed.group(2).strip()
        elif '——' in heading:
            layer, name = heading.split('——', 1)
            name = name.strip()
            identity = re.sub(r'^[一二三]层(?:补充)?[：:]?\s*', '', layer).strip() or '重要角色'
        elif re.match(r'^[1-9]\d*\.\s*', heading):
            name = re.sub(r'^[1-9]\d*\.\s*', '', heading).strip()
            identity = '重要角色'
        if not name:
            continue
        if name in seen:
            continue
        seen.add(name)
        outer = find_field(section, 'outer') or '档案以其公开身份和行动表现为主要外貌呈现。'
        personality = find_field(section, 'personality') or '通过行为、选择和语言特征逐步呈现。'
        roles.append((name, identity, outer, personality))
    return roles


def card_lines(state):
    paywall = state.get('paywallEpisodes') or state.get('paywall') or []
    labels = ['一卡（前段）', '二卡（中段）', '三卡（终段）']
    groups = []
    if paywall:
        third = max(1, (len(paywall) + 2) // 3)
        groups = [paywall[i:i + third] for i in range(0, len(paywall), third)]
    while len(groups) < 3:
        groups.append([])
    return [
        f'{label}：' + ('、'.join(f'第{n}集' for n in group) if group else '待定')
        for label, group in zip(labels, groups)
    ]


def episode_files():
    return sorted(EPISODES.glob('ep*.md'), key=lambda path: int(path.stem[-3:]))


def episode_title_and_body(path):
    lines = path.read_text(encoding='utf-8').splitlines()
    title = lines[0].lstrip('#').strip() if lines else path.stem
    return title, '\n'.join(lines[1:]).strip()


def clean_brief(text):
    """只清理格式，不裁剪正文；简纲完整性必须保留。"""
    text = re.sub(r'^#{1,6}\s*简纲\s*$', '', text, flags=re.M)
    text = re.sub(r'\*\*', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    return text


def validate_brief(brief):
    """检测简纲长度，超限时停止导出，绝不静默截断。"""
    count = len(brief.replace('\n', '').replace('\r', ''))
    if count > 300:
        raise ValueError(
            f'简纲共{count}字，超过300字限制。请压缩简纲内容后再导出；脚本不会自动截断，避免故事不完整。'
        )
    return count


def brief_from_episode_directory():
    """Turn the act/episode beats into a compact synopsis for the export."""
    path = ROOT / 'episode-directory.md'
    if not path.exists():
        return ''
    text = path.read_text(encoding='utf-8')
    explicit = re.search(r'^##\s*(?:简纲|故事梗概|故事简介)\s*$\n(.*?)(?=^##\s|\Z)', text, flags=re.M | re.S)
    if explicit:
        return clean_brief(explicit.group(1))

    # 分集目录末尾的节奏检查是全剧主线摘要，优先使用它，避免只拼接前几集而
    # 把中后段和结局截掉。再补入最终集，确保简纲包含结局落点。
    pace_matches = re.findall(
        r'^-\s+\*\*(起势段|攀升段|风暴段|决战段)：\*\*\s*(.+)$',
        text,
        flags=re.M,
    )
    if pace_matches:
        summary = '；'.join(f'{label}：{content.rstrip("。！？；")}' for label, content in pace_matches)
        ending = re.search(
            r'^第72集：\s*(.*?)\s*——\s*(.*?)(?:\s+\[.*?\])?(?:\s+[🔥💰]+)?\s*$',
            text,
            flags=re.M,
        )
        if ending:
            summary += f'；结局：{ending.group(2).rstrip("。！？；")}。'
        return clean_brief(summary)

    blocks = re.findall(r'^##\s+([^\n]*幕[^\n]*)\n(.*?)(?=^##\s|\Z)', text, flags=re.M | re.S)
    summaries = []
    episode_pattern = re.compile(r'^第(\d+)集：\s*(.*?)\s*——\s*(.*?)(?:\s+\[.*?\])?(?:\s+[🔥💰]+)?\s*$')
    for act, block in blocks:
        entries = []
        for line in block.splitlines():
            match = episode_pattern.match(line.strip())
            if match:
                entries.append((int(match.group(1)), match.group(2), match.group(3)))
        if not entries:
            continue
        first = entries[0]
        last = entries[-1]
        # Prefer marked beats, but keep the synopsis compact enough for front matter.
        marked = []
        for line in block.splitlines():
            match = episode_pattern.match(line.strip())
            if match and any(mark in line for mark in ('🔥', '💰')):
                item = (int(match.group(1)), match.group(2), match.group(3))
                if item[0] not in (first[0], last[0]):
                    marked.append(item)
        key = entries[1:-1]
        chosen = marked[:3] if marked else key[:2]
        trim = lambda value: value.rstrip('。！？；')
        middle = '；'.join(f'第{n}集“{title}”：{trim(synopsis)}' for n, title, synopsis in chosen)
        summary = f'{act}：从“{first[1]}”切入，{trim(first[2])}。'
        if middle:
            summary += f'中段关键转折包括{middle}。'
        summary += f'最终以“{last[1]}”收束：{trim(last[2])}。'
        summaries.append(summary)
    return '\n\n'.join(summaries)


def extract_brief(state):
    """Read an explicit brief first, then derive one from the story plan."""
    for path in (ROOT / 'story-brief.md', ROOT / 'brief.md'):
        if path.exists() and path.read_text(encoding='utf-8').strip():
            return clean_brief(path.read_text(encoding='utf-8'))
    for key in ('storyBrief', 'brief', 'synopsis'):
        if state.get(key):
            return clean_brief(str(state[key]))
    directory_brief = brief_from_episode_directory()
    if directory_brief:
        return clean_brief(directory_brief)
    plan = ROOT / 'creative-plan.md'
    if plan.exists():
        text = plan.read_text(encoding='utf-8')
        for pattern in (
            r'^##\s*(?:简纲|故事梗概|故事简介)\s*$\n(.*?)(?=^##\s|\Z)',
            r'^###\s*主线[^\n]*\n(.*?)(?=^###\s|^##\s|\Z)',
        ):
            match = re.search(pattern, text, flags=re.M | re.S)
            if match:
                brief = clean_brief(match.group(1))
                if brief:
                    return brief
    return ''


def build_markdown(state, roles, files, brief):
    title = state.get('dramaTitle') or '未命名短剧'
    total = int(state.get('totalEpisodes') or len(files) or 0)
    completed = [int(path.stem[-3:]) for path in files]
    genres = state.get('genre') or []
    if isinstance(genres, str):
        genres = [part.strip() for part in re.split(r'[+、,，]', genres) if part.strip()]
    genre_label = ' + '.join(genres) if genres else '未分类'
    mode = state.get('mode') or ''
    if mode == 'domestic':
        mode_label = '国内标准格式'
    elif mode == 'hollywood':
        mode_label = '好莱坞行业标准格式'
    elif mode:
        mode_label = str(mode)
    else:
        mode_label = '未指定'
    lines = [
        f'# {title}', '',
        f'> 题材：{genre_label}｜输出模式：{mode_label}｜创作日期：{date.today().isoformat()}',
        f'> 当前进度：已完成 {len(completed)}/{total} 集｜阶段导出，未完成集数以占位标记保留。', '',
        '## 卡点', '',
    ]
    for line in card_lines(state):
        lines += [line, '']
    lines += ['## 角色', '']
    for name, identity, outer, personality in roles:
        lines += [f'### {name}（{identity}）', '', f'**外貌：** {outer}', '', f'**性格：** {personality}', '']
    if brief:
        lines += ['## 简纲', '', brief, '']
    lines += ['## 分集剧本', '']
    for path in files:
        title_line, body = episode_title_and_body(path)
        lines += [f'### {title_line}', '', body, '']
    for number in range(1, total + 1):
        if number not in completed:
            lines += [f'### 第{number}集：（待撰写）', '', '本集剧本待撰写。', '']
    return '\n'.join(lines).rstrip() + '\n'


def configure_styles(doc):
    styles = doc.styles
    set_style_font(styles['Normal'], size=11)
    set_spacing(styles['Normal'], 0, 8, 1.333, WD_ALIGN_PARAGRAPH.JUSTIFY)
    for style_name, size, color, before, after in [
        ('Title', 30, NAVY, 0, 8), ('Subtitle', 14, DARK_BLUE, 0, 8),
        ('Heading 1', 16, BLUE, 18, 10), ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        set_style_font(style, size=size, color=color, bold=style_name.startswith('Heading'))
        set_spacing(style, before, after, 1.15, WD_ALIGN_PARAGRAPH.LEFT)
        style.paragraph_format.keep_with_next = True

    def make_style(name, size=11, color=RGBColor(0, 0, 0), bold=False, italic=False, before=0, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size=size, color=color, bold=bold)
        style.font.italic = italic
        set_spacing(style, before, after, 1.25, align)

    make_style('CoverKicker', 11, GOLD, True, False, 0, 18, WD_ALIGN_PARAGRAPH.CENTER)
    make_style('CoverMeta', 10.5, GRAY, False, False, 0, 5, WD_ALIGN_PARAGRAPH.CENTER)
    make_style('SceneHeader', 11, DARK_BLUE, True, False, 8, 3)
    make_style('CastLine', 10.5, GRAY, False, False, 0, 5)
    make_style('ActionLine', 10.5, RGBColor(60, 60, 60), False, True, 0, 5, WD_ALIGN_PARAGRAPH.JUSTIFY)
    make_style('SpecialLine', 10.5, GOLD, True, False, 4, 5)
    make_style('HookLine', 10.5, DARK_BLUE, False, False, 4, 4)
    make_style('Placeholder', 10.5, GRAY, False, False, 3, 2)


def set_reference_paragraph(paragraph, text=''):
    """Apply the retained reference's plain screenplay paragraph treatment."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 2.0
    ppr = paragraph._p.get_or_add_pPr()
    rpr = ppr.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        ppr.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for key in ('hint', 'ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{key}'), FONT)
    for tag in ('sz', 'szCs'):
        size = rpr.find(qn(f'w:{tag}'))
        if size is None:
            size = OxmlElement(f'w:{tag}')
            rpr.append(size)
        size.set(qn('w:val'), '21')
    if text:
        run = paragraph.add_run(text)
        set_font(run, name=FONT, size=10.5)
    return paragraph


def chinese_number(number):
    digits = '零一二三四五六七八九'
    if number < 10:
        return digits[number]
    if number < 20:
        return '十' if number == 10 else '十' + digits[number - 10]
    if number < 100:
        tens, ones = divmod(number, 10)
        return digits[tens] + '十' + (digits[ones] if ones else '')
    return str(number)


def build_docx(state, roles, files, brief):
    title = state.get('dramaTitle') or '未命名短剧'
    total = int(state.get('totalEpisodes') or len(files) or 0)
    completed = {int(path.stem[-3:]) for path in files}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(11906)
    section.page_height = Twips(16838)
    section.left_margin = section.right_margin = Twips(1800)
    section.top_margin = section.bottom_margin = Twips(1440)
    section.header_distance = Twips(851)
    section.footer_distance = Twips(992)

    # The retained reference has no visible header/footer or decorative cover.
    add = lambda text='': set_reference_paragraph(doc.add_paragraph(), text)
    add(f'《{title}》')
    add('卡点：')
    for line in card_lines(state):
        add(line)
    add()

    for name, identity, outer, personality in roles:
        add(f'{name}（{identity}）')
        add(f'外貌：{outer}')
        add(f'性格：{personality}')
        add()

    if brief:
        add('简纲：')
        for paragraph in re.split(r'\n\s*\n', brief):
            add(paragraph.strip())
        add()

    for path in files:
        number = int(path.stem[-3:])
        _, body = episode_title_and_body(path)
        add(f'第{chinese_number(number)}集')
        for raw in body.splitlines():
            line = raw.strip()
            if line.startswith('>'):
                continue
            line = line.replace('**', '')
            add(line)

    # Keep the same plain continuous stream for unfinished episodes.
    for number in range(1, total + 1):
        if number not in completed:
            add(f'第{chinese_number(number)}集')
            add('本集剧本待撰写。')

    output = EXPORT / f'{title}-完整剧本.docx'
    doc.save(output)
    return output


def main():
    if not EPISODES.exists() or not list(EPISODES.glob('ep*.md')):
        raise FileNotFoundError(f'未找到分集剧本目录或文件：{EPISODES}')
    EXPORT.mkdir(exist_ok=True)
    state = json.loads(STATE_PATH.read_text(encoding='utf-8'))
    title = state.get('dramaTitle') or '未命名短剧'
    roles = parse_roles()
    files = episode_files()
    brief = extract_brief(state)
    brief_count = validate_brief(brief) if brief else 0
    markdown = build_markdown(state, roles, files, brief)
    md_path = EXPORT / f'{title}-完整剧本.md'
    md_path.write_text(markdown, encoding='utf-8')
    docx_path = build_docx(state, roles, files, brief)
    print(md_path)
    print(docx_path)
    print(f'roles={len(roles)} episodes={len(files)} brief_chars={brief_count}')


if __name__ == '__main__':
    main()
